"""
CramCoach - Exam Panic Mode Edition
=====================================
A Streamlit + Gemini app that turns EITHER a rambled voice note OR a
photo of handwritten/printed notes into a triaged study guide, a
flashcard deck, and a difficulty-tagged quiz - graded locally, in real
time, against how many hours you have left until your exam.

Author: <your name>
Capstone Category: B - EdTech & Campus Survival (#8, customized)
"""

import io
import json
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from google import genai
from google.genai import types

# --------------------------------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="CramCoach | Exam Panic Mode",
    page_icon="⏰",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --------------------------------------------------------------------------
# CONSTANTS
# --------------------------------------------------------------------------
DIFFICULTIES = {
    "Easy": "recall-level facts and definitions, gentle pacing",
    "Medium": "applied understanding, some multi-step reasoning",
    "Hard": "exam-style tricky questions, edge cases, comparisons",
    "Panic Mode": "only the highest-yield, most-likely-to-be-tested concepts, "
    "ruthlessly prioritized, minimal fluff",
}

MAX_DOCX_CHARS = 15000  # guardrail so extracted text stays within a safe prompt size

SYSTEM_PROMPT = """You are "The Cram Coach" - a blunt, funny, high-energy
exam-prep TA who has pulled a hundred all-nighters and has zero patience
for procrastination, but genuinely wants the student to pass. Your humor
is urgency-driven ("you have X hours, let's move") never mean-spirited.

You will receive ONE of these inputs from the student:
(A) An audio recording of them rambling their lecture notes out loud,
(B) A photo of their handwritten or printed notes / textbook page,
(C) A PDF document (lecture slides, scanned notes, or a textbook excerpt), OR
(D) Extracted text from a Word document they uploaded.

For audio: transcribe it mentally and triage the content.
For a photo or PDF: read every word you can make out - including messy
handwriting - and triage the content. If parts are illegible, do your
best to infer from context and note briefly in the study guide that a
section was hard to read, rather than inventing facts.
For extracted document text: treat it as ground truth and triage it
directly.

Either way, convert the input into exam-ready study material.

You must respond with STRICT, VALID JSON ONLY. No markdown fences, no
preamble, no trailing commentary. Match this exact schema:

{
  "topic_detected": "<short topic name inferred from the input>",
  "pep_talk": "<one punchy, urgency-flavored coach line, max 25 words>",
  "study_guide": ["<high-yield bullet point 1>", "<bullet 2>", "..."],
  "flashcards": [
    {"front": "<question or term>", "back": "<answer or definition>"}
  ],
  "quiz": [
    {
      "question": "<quiz question text>",
      "options": ["<option A>", "<option B>", "<option C>", "<option D>"],
      "correct_index": <integer 0-3>,
      "explanation": "<why this answer is correct, 1-2 sentences>",
      "difficulty": "<Easy|Medium|Hard>"
    }
  ]
}

Generate between 5 and 8 flashcards and exactly 5 quiz questions. Tailor
depth and question difficulty to the requested difficulty level and the
time pressure described in the prompt.
"""


# --------------------------------------------------------------------------
# SESSION STATE INIT
# --------------------------------------------------------------------------
def init_state():
    defaults = {
        "history": [],           # cram sessions: topic, difficulty, score, timestamp
        "last_result": None,     # last Gemini JSON result
        "last_source": None,     # "🎙️ Voice" or "📸 Photo"
        "quiz_answers": {},      # user's selected option indices per question
        "quiz_graded": False,    # whether current quiz has been graded
        "flashcards_df": None,   # editable flashcard deck
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


init_state()

# --------------------------------------------------------------------------
# SIDEBAR - CONFIG
# --------------------------------------------------------------------------
with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    api_key = st.text_input(
        "Gemini API Key",
        type="password",
        help="Get a free key at https://aistudio.google.com/apikey. "
        "Never hardcoded in the repo.",
    )
    st.caption("🔒 Held only in this browser session's memory.")

    st.divider()
    st.markdown("## ⏱️ Exam Pressure")
    hours_left = st.number_input(
        "Hours until your exam", min_value=1, max_value=72, value=6, step=1
    )
    difficulty = st.selectbox("Difficulty", options=list(DIFFICULTIES.keys()), index=1)

    if hours_left <= 3:
        st.error(f"🚨 {hours_left}h left. PANIC MODE ENGAGED.")
    elif hours_left <= 12:
        st.warning(f"⚠️ {hours_left}h left. Let's move.")
    else:
        st.info(f"🕐 {hours_left}h left. Still time to breathe.")

    st.divider()
    st.markdown("## 📊 Readiness Stats")
    if st.session_state.history:
        avg_score = sum(h["Score"] for h in st.session_state.history) / len(
            st.session_state.history
        )
        st.metric("Average Quiz Score", f"{avg_score:.0f}%")
        st.metric("Cram Sessions", len(st.session_state.history))
    else:
        st.caption("No sessions yet. Record your first ramble to begin.")

    st.divider()
    st.caption("Built for MirAI Capstone — Category B #8 (Customized)")

# --------------------------------------------------------------------------
# HEADER
# --------------------------------------------------------------------------
st.title("⏰ CramCoach")
st.markdown(
    "##### Ramble your notes, snap a photo, or upload a PDF/DOCX/JPG. Get "
    "a triaged study guide, flashcards, and an auto-graded quiz — before "
    "the clock runs out."
)
st.divider()

# --------------------------------------------------------------------------
# GEMINI CALL (native audio / image / PDF input, or extracted DOCX text -
# one shared pipeline across all four modalities)
# --------------------------------------------------------------------------
def generate_cram_kit(
    difficulty: str,
    hours_left: int,
    input_kind: str,  # "audio" | "photo" | "pdf" | "docx_text"
    input_bytes: bytes = None,
    mime_type: str = None,
    extracted_text: str = None,
) -> dict:
    """Send raw audio, a photo, or a PDF directly to Gemini's multimodal
    endpoint (native understanding - no separate transcription/OCR step),
    OR pass pre-extracted DOCX text as plain context. Either way, returns
    the same structured JSON study kit."""
    client = genai.Client(api_key=api_key)

    difficulty_desc = DIFFICULTIES[difficulty]

    instruction_lines = {
        "audio": "Listen to the attached audio of the student's rambled "
        "lecture summary.",
        "photo": "Read the attached photo of the student's handwritten or "
        "printed notes, including any diagrams, labels, or margin "
        "annotations you can make out.",
        "pdf": "Read the attached PDF document, which may contain lecture "
        "slides, scanned notes, or a textbook excerpt.",
        "docx_text": "Read the following text extracted from the student's "
        "uploaded Word document:\n\n" + (extracted_text or ""),
    }
    instruction_line = instruction_lines[input_kind]

    user_prompt = f"""
    The student has {hours_left} hour(s) until their exam.
    Requested difficulty: {difficulty} ({difficulty_desc}).

    {instruction_line}
    Triage it into the exam-ready JSON structure defined in the system
    prompt. Prioritize whatever would yield the most exam points given
    the time remaining.
    """

    # DOCX path: plain text only, no binary Part needed.
    if input_kind == "docx_text":
        contents = [user_prompt]
    else:
        contents = [
            types.Part.from_bytes(data=input_bytes, mime_type=mime_type),
            user_prompt,
        ]

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=contents,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=0.8,
            response_mime_type="application/json",
        ),
    )

    raw = response.text.strip().replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract plain text from an uploaded .docx file using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text out of any tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs)
    if len(text) > MAX_DOCX_CHARS:
        text = text[:MAX_DOCX_CHARS] + "\n[...truncated for length...]"
    return text


# --------------------------------------------------------------------------
# CAPTURE - two input modes, same downstream pipeline
# --------------------------------------------------------------------------
st.subheader("1️⃣ Feed the Coach Your Notes")

input_mode = st.radio(
    "Choose your input method",
    options=["🎙️ Voice Ramble", "📸 Photo of Notes", "📄 Upload File (PDF / DOCX / JPG)"],
    horizontal=True,
)

with st.form("form_capture", clear_on_submit=False):
    audio = photo = uploaded_file = None

    if input_mode == "🎙️ Voice Ramble":
        audio = st.audio_input("Record yourself explaining what you remember (30-90 sec is ideal)")
    elif input_mode == "📸 Photo of Notes":
        photo = st.camera_input("Snap a photo of your handwritten or printed notes")
    else:
        uploaded_file = st.file_uploader(
            "Upload your notes",
            type=["pdf", "docx", "jpg", "jpeg"],
            accept_multiple_files=False,
        )

    submitted = st.form_submit_button(
        "🔥 Cram It", use_container_width=True, type="primary"
    )

if submitted:
    if not api_key:
        st.error("⚠️ Please enter your Gemini API key in the sidebar first.")
    elif input_mode == "🎙️ Voice Ramble" and not audio:
        st.warning("🎙️ Please record a voice note before submitting.")
    elif input_mode == "📸 Photo of Notes" and not photo:
        st.warning("📸 Please capture a photo of your notes before submitting.")
    elif input_mode == "📄 Upload File (PDF / DOCX / JPG)" and not uploaded_file:
        st.warning("📄 Please upload a PDF, DOCX, or JPG file before submitting.")
    else:
        with st.spinner("Coach is triaging your notes..."):
            try:
                if input_mode == "🎙️ Voice Ramble":
                    result = generate_cram_kit(
                        difficulty, hours_left, "audio",
                        input_bytes=audio.getvalue(), mime_type="audio/wav",
                    )
                    source_label = "🎙️ Voice"

                elif input_mode == "📸 Photo of Notes":
                    result = generate_cram_kit(
                        difficulty, hours_left, "photo",
                        input_bytes=photo.getvalue(), mime_type="image/jpeg",
                    )
                    source_label = "📸 Photo"

                else:
                    file_bytes = uploaded_file.getvalue()
                    ext = uploaded_file.name.rsplit(".", 1)[-1].lower()

                    if ext == "pdf":
                        result = generate_cram_kit(
                            difficulty, hours_left, "pdf",
                            input_bytes=file_bytes, mime_type="application/pdf",
                        )
                        source_label = "📄 PDF"

                    elif ext == "docx":
                        extracted = extract_docx_text(file_bytes)
                        if not extracted.strip():
                            st.warning(
                                "Couldn't find readable text in that DOCX — "
                                "is it empty or image-only?"
                            )
                            st.stop()
                        result = generate_cram_kit(
                            difficulty, hours_left, "docx_text",
                            extracted_text=extracted,
                        )
                        source_label = "📄 DOCX"

                    else:  # jpg / jpeg
                        result = generate_cram_kit(
                            difficulty, hours_left, "photo",
                            input_bytes=file_bytes, mime_type="image/jpeg",
                        )
                        source_label = "📄 Photo File"

                st.session_state.last_result = result
                st.session_state.last_source = source_label
                st.session_state.quiz_answers = {}
                st.session_state.quiz_graded = False
                st.session_state.flashcards_df = pd.DataFrame(result.get("flashcards", []))
            except json.JSONDecodeError:
                st.error("The AI response couldn't be parsed. Please try again.")
            except Exception as e:
                st.error(f"Something went wrong: {e}")

# --------------------------------------------------------------------------
# RESULTS
# --------------------------------------------------------------------------
result = st.session_state.last_result

if result:
    st.divider()
    st.subheader("2️⃣ Your Cram Kit")

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Topic Detected", result.get("topic_detected", "—"))
    m2.metric("Source", st.session_state.last_source or "—")
    m3.metric("Flashcards Generated", len(result.get("flashcards", [])))
    m4.metric("Hours Left", hours_left)

    st.warning(f"**\"{result.get('pep_talk', '')}\"** — Cram Coach")

    col_guide, col_cards = st.columns([1, 1.2], gap="large")

    with col_guide:
        with st.expander("📋 Triaged Study Guide", expanded=True):
            for point in result.get("study_guide", []):
                st.markdown(f"- {point}")

    with col_cards:
        with st.expander("🗂️ Flashcard Deck (editable)", expanded=True):
            if st.session_state.flashcards_df is not None:
                edited = st.data_editor(
                    st.session_state.flashcards_df,
                    use_container_width=True,
                    num_rows="dynamic",
                    hide_index=True,
                    column_config={
                        "front": st.column_config.TextColumn("Front", width="medium"),
                        "back": st.column_config.TextColumn("Back", width="large"),
                    },
                )
                st.session_state.flashcards_df = edited

    # ----------------------------------------------------------------
    # QUIZ - locally auto-graded, difficulty-tagged
    # ----------------------------------------------------------------
    st.divider()
    st.subheader("3️⃣ Prove It — Quiz Time")

    quiz = result.get("quiz", [])
    if quiz:
        with st.form("quiz_form"):
            for i, q in enumerate(quiz):
                st.markdown(
                    f"**Q{i + 1}. {q['question']}**  "
                    f"`{q.get('difficulty', difficulty)}`"
                )
                choice = st.radio(
                    f"quiz_q_{i}",
                    options=list(range(len(q["options"]))),
                    format_func=lambda idx, opts=q["options"]: opts[idx],
                    key=f"quiz_radio_{i}",
                    label_visibility="collapsed",
                )
                st.session_state.quiz_answers[i] = choice
                st.markdown("---")

            grade_submitted = st.form_submit_button(
                "✅ Grade Me", use_container_width=True, type="primary"
            )

        if grade_submitted:
            st.session_state.quiz_graded = True

        if st.session_state.quiz_graded:
            correct_count = 0
            weak_topics = []

            for i, q in enumerate(quiz):
                user_choice = st.session_state.quiz_answers.get(i)
                is_correct = user_choice == q["correct_index"]
                if is_correct:
                    correct_count += 1
                else:
                    weak_topics.append(q.get("difficulty", "Unknown"))

                icon = "✅" if is_correct else "❌"
                with st.expander(f"{icon} Q{i + 1}: {q['question'][:60]}..."):
                    st.write(f"Your answer: {q['options'][user_choice]}")
                    st.write(f"Correct answer: {q['options'][q['correct_index']]}")
                    st.caption(q.get("explanation", ""))

            score_pct = round((correct_count / len(quiz)) * 100)

            prev_scores = [h["Score"] for h in st.session_state.history]
            delta = score_pct - prev_scores[-1] if prev_scores else None

            r1, r2, r3 = st.columns(3)
            r1.metric(
                "Quiz Score",
                f"{score_pct}%",
                delta=f"{delta:+d}%" if delta is not None else None,
            )
            r2.metric("Correct", f"{correct_count}/{len(quiz)}")
            weakest = max(set(weak_topics), key=weak_topics.count) if weak_topics else "None"
            r3.metric("Weakest Tier", weakest)

            if score_pct >= 80:
                st.success("🎉 You're actually ready. Go eat something and sleep.")
            elif score_pct >= 50:
                st.warning("😬 Getting there. Review the ❌ items above, then re-cram.")
            else:
                st.error("🚨 Not ready yet. Re-record a more detailed ramble and try again.")

            # Log this cram session to history
            already_logged = any(
                h.get("Topic") == result.get("topic_detected")
                and h.get("_run") == id(result)
                for h in st.session_state.history
            )
            if not already_logged:
                st.session_state.history.append(
                    {
                        "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "Topic": result.get("topic_detected", "—"),
                        "Source": st.session_state.last_source or "—",
                        "Difficulty": difficulty,
                        "Score": score_pct,
                        "_run": id(result),
                    }
                )
else:
    st.caption("Your study guide, flashcards, and quiz will appear here after you submit your notes.")

# --------------------------------------------------------------------------
# SESSION HISTORY
# --------------------------------------------------------------------------
st.divider()
st.subheader("📈 Cram History")

if st.session_state.history:
    df = pd.DataFrame(st.session_state.history).drop(columns=["_run"], errors="ignore")
    st.data_editor(
        df,
        width="stretch",
        hide_index=True,
        column_config={
            "Score": st.column_config.ProgressColumn(
                "Score", min_value=0, max_value=100, format="%d%%"
            ),
        },
    )
    st.line_chart(df.set_index("Timestamp")["Score"])
else:
    st.caption("No cram sessions logged yet.")

st.divider()
st.caption(
    "⚠️ Disclaimer: AI-generated study material for exam-prep support only. "
    "Always verify against your actual course materials before an exam."
)
